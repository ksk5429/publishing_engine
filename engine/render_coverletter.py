"""
Generate a cover letter DOCX for journal submission.

Reads cover letter content from a cover_letter.md or YAML metadata.

Usage:
    python _shared/render_coverletter.py _shared/sample
"""
from __future__ import annotations
import argparse, logging, sys, re
from pathlib import Path
from datetime import date
sys.path.insert(0, str(Path(__file__).parent))
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Inches, RGBColor
from qmd_parser import parse_qmd

PAPERS_DIR = Path(__file__).parent.parent  # repo root
BODY_FONT = "Times New Roman"
NAVY = RGBColor(0x00, 0x33, 0x66)
logging.basicConfig(level=logging.INFO, format="[%(asctime)s] %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("coverletter")

SAMPLE_LETTER = {
    "to": "Editor, Computers and Geotechnics",
    "subject": "Submission of manuscript for review",
    "body": [
        "We submit the enclosed manuscript entitled \"{title}\" for consideration as an original research article in {journal}.",
        "This paper addresses {gap}. The key contributions are: {contributions}.",
        "All authors have approved the manuscript. The work has not been published elsewhere and is not under consideration at another journal.",
        "We confirm that generative AI (Anthropic Claude) was used for academic writing coaching and script scaffolding only. All numerical values were independently computed. The authors take full responsibility for the publication.",
        "We look forward to hearing from you.",
    ],
    "gap": "the open question of which vibration-derived feature reliably detects scour on tripod suction-bucket foundations across multiple soil conditions",
    "contributions": "(1) systematic 64-feature evaluation identifying strain-acceleration coherence as the only sign-consistent feature, (2) a methodological correction preventing propagation of the incorrect fixity-ratio claim, and (3) a definitive bound on cross-site transfer feasibility",
    "journal": "Computers and Geotechnics",
}


def render_coverletter(paper_name: str) -> Path:
    paper_dir = PAPERS_DIR / paper_name
    meta, _ = parse_qmd(paper_dir / "manuscript.qmd")
    title = meta.get("title", paper_name)

    cl_path = paper_dir / "cover_letter.md"
    if cl_path.exists():
        cl_text = cl_path.read_text(encoding="utf-8")
        # Parse YAML frontmatter from cover letter
        if cl_text.startswith("---"):
            end = cl_text.find("---", 3)
            import yaml
            cl_meta = yaml.safe_load(cl_text[3:end]) or {}
            cl_body = cl_text[end+3:].strip()
        else:
            cl_meta = {}
            cl_body = cl_text.strip()
    else:
        cl_meta = SAMPLE_LETTER
        cl_body = None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Date
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(date.today().strftime("%B %d, %Y"))
    r.font.name = BODY_FONT
    r.font.size = Pt(11)

    doc.add_paragraph()

    # Addressee
    to = cl_meta.get("to", "Dear Editor,")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(to)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)

    doc.add_paragraph()

    # Subject
    subject = cl_meta.get("subject", f"Re: {title}")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(f"Re: {title}")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.bold = True

    doc.add_paragraph()

    # Body
    if cl_body:
        for para_text in cl_body.split("\n\n"):
            para_text = para_text.strip()
            if not para_text or para_text.startswith("---"):
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(para_text)
            r.font.name = BODY_FONT
            r.font.size = Pt(11)
    else:
        journal = cl_meta.get("journal", "the journal")
        gap = cl_meta.get("gap", "an important research question")
        contribs = cl_meta.get("contributions", "several novel contributions")
        body_paras = cl_meta.get("body", SAMPLE_LETTER["body"])
        for para_template in body_paras:
            text = para_template.format(
                title=title, journal=journal, gap=gap, contributions=contribs,
            )
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(text)
            r.font.name = BODY_FONT
            r.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("Sincerely,")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("Kyeong-Sun Kim")
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("Department of Civil and Environmental Engineering\n"
                   "Seoul National University\n"
                   "kyeongsunkim@snu.ac.kr")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    out_dir = paper_dir / "_output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "cover_letter.docx"
    doc.save(str(out_path))
    log.info("wrote %s (%d bytes)", out_path, out_path.stat().st_size)
    return out_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("paper")
    args = parser.parse_args()
    render_coverletter(args.paper)

if __name__ == "__main__":
    main()
