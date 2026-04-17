"""
Shared DOCX rendering engine for all PhD papers.

Replicates Paper A's publication-grade styling with full control over
every element. No Quarto dependency — builds DOCX from scratch using
python-docx.

Features:
  - Times New Roman 11pt body, 14/12/11pt headings (navy #003366)
  - Proper title page with author, affiliation, email, ORCID
  - Abstract block (italic, reduced margins)
  - Keywords and Highlights
  - Section numbering (1, 1.1, 2, 2.1, ...)
  - Figures embedded with numbered captions
  - Booktabs tables (header shading, no vertical borders)
  - Equations rendered as PNG via matplotlib mathtext
  - Running header + page numbers in footer
  - CRediT, COI, AI disclosure, Acknowledgements blocks
  - Hanging-indent reference list

Usage by per-paper renderers:
    from _shared.docx_engine import DocxBuilder
    b = DocxBuilder(title="...", authors=[...], running_header="...")
    b.heading("Introduction", level=1)
    b.paragraph("Body text...")
    b.figure("figures/fig1.png", "Fig. 1", "Caption text.")
    b.save("manuscript_styled.docx")
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from tempfile import mkdtemp

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)
H3_SIZE = Pt(11)
CAP_SIZE = Pt(10)
SMALL_SIZE = Pt(9)

NAVY = RGBColor(0x00, 0x33, 0x66)
DARK_NAVY = RGBColor(0x1F, 0x38, 0x64)
INK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "D9E1F2"
ALT_ROW_FILL = "F5F7FB"

log = logging.getLogger("docx_engine")


def _shade(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _remove_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    if tbl_pr is None:
        return
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _add_rule(row, edge="bottom", color="003366", sz="6"):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
        tc_pr.append(borders)


def _page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


class DocxBuilder:
    def __init__(
        self,
        title: str,
        authors: list[dict],
        running_header: str,
        abstract: str = "",
        keywords: list[str] | None = None,
        highlights: list[str] | None = None,
    ):
        self.doc = Document()
        self.title = title
        self.authors = authors
        self.running_header = running_header
        self.abstract_text = abstract
        self.keywords = keywords or []
        self.highlights = highlights or []
        self._section_counters = [0, 0, 0]
        self._fig_num = 0
        self._tbl_num = 0
        self._eq_num = 0
        self._eq_tmp = Path(mkdtemp(prefix="eq_"))
        self._configure()
        self._title_page()

    def _configure(self):
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        style = self.doc.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.first_line_indent = Inches(0.25)
        style.paragraph_format.space_after = Pt(4)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.running_header)
        r.font.name = BODY_FONT
        r.font.size = SMALL_SIZE
        r.font.italic = True
        r.font.color.rgb = GREY

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _page_number(fp)

    def _title_page(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(48)
        r = p.add_run(self.title)
        r.font.name = BODY_FONT
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = NAVY

        self.doc.add_paragraph()  # spacer

        # Build affiliation map: unique affiliations → letter labels (a, b, c...)
        affil_map: dict[str, str] = {}
        label_idx = 0
        for au in self.authors:
            aff = au.get("affiliation", "")
            if aff and aff not in affil_map:
                affil_map[aff] = chr(ord("a") + label_idx)
                label_idx += 1

        # Author names in single line: "Name^a, Name^a,*"
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        for i, au in enumerate(self.authors):
            if i > 0:
                r = p.add_run(", ")
                r.font.name = BODY_FONT
                r.font.size = Pt(11)
            r = p.add_run(au["name"])
            r.font.name = BODY_FONT
            r.font.size = Pt(11)
            r.font.bold = True
            # Superscript affiliation label
            aff = au.get("affiliation", "")
            label = affil_map.get(aff, "")
            sup_text = label
            if au.get("corresponding"):
                sup_text += ",*" if label else "*"
            if sup_text:
                r2 = p.add_run(f" {sup_text}")
                r2.font.name = BODY_FONT
                r2.font.size = Pt(8)
                r2.font.superscript = True

        # Affiliation lines with letter labels
        for aff, label in affil_map.items():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{label} ")
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.superscript = True
            r2 = p.add_run(aff)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)
            r2.font.italic = True

        # Corresponding author email
        corr = [a for a in self.authors if a.get("corresponding")]
        if corr:
            email = corr[0].get("email", "")
            if email:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(4)
                r = p.add_run("* Corresponding author: ")
                r.font.name = BODY_FONT
                r.font.size = SMALL_SIZE
                r.font.color.rgb = GREY
                r2 = p.add_run(email)
                r2.font.name = BODY_FONT
                r2.font.size = SMALL_SIZE
                r2.font.color.rgb = NAVY

        # ── Abstract block (italic, indented, reduced size) ──
        if self.abstract_text:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("Abstract")
            r.font.name = BODY_FONT
            r.font.size = H2_SIZE
            r.font.bold = True
            r.font.color.rgb = NAVY

            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(self.abstract_text.strip())
            r.font.name = BODY_FONT
            r.font.size = Pt(10)
            r.font.italic = True

        # ── Keywords ──
        if self.keywords:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run("Keywords: ")
            r.font.name = BODY_FONT
            r.font.size = Pt(10)
            r.font.bold = True
            r2 = p.add_run("; ".join(self.keywords))
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)

        # ── Highlights ──
        if self.highlights:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run("Highlights")
            r.font.name = BODY_FONT
            r.font.size = H2_SIZE
            r.font.bold = True
            r.font.color.rgb = NAVY
            for h in self.highlights:
                p = self.doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(f"\u2022  {h}")
                r.font.name = BODY_FONT
                r.font.size = Pt(10)

        self.doc.add_page_break()

        # ── Table of Contents field (user updates in Word via F9) ──
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run("Table of Contents")
        r.font.name = BODY_FONT
        r.font.size = H1_SIZE
        r.font.bold = True
        r.font.color.rgb = NAVY

        toc_para = self.doc.add_paragraph()
        toc_para.paragraph_format.first_line_indent = Pt(0)
        r = toc_para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-3" \h \z \u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder_r = OxmlElement("w:r")
        placeholder_t = OxmlElement("w:t")
        placeholder_t.text = "Right-click → Update Field to populate (Word: References → Update Table)."
        placeholder_r.append(placeholder_t)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_sep)
        toc_para._p.append(placeholder_r)
        r._r.append(fld_end)

        self.doc.add_page_break()

    def heading(self, text: str, level: int = 1, numbered: bool = True) -> str:
        if numbered:
            if level == 1:
                self._section_counters[0] += 1
                self._section_counters[1] = 0
                self._section_counters[2] = 0
                num = f"{self._section_counters[0]}. "
            elif level == 2:
                self._section_counters[1] += 1
                self._section_counters[2] = 0
                num = f"{self._section_counters[0]}.{self._section_counters[1]}. "
            else:
                self._section_counters[2] += 1
                num = f"{self._section_counters[0]}.{self._section_counters[1]}.{self._section_counters[2]}. "
            display = f"{num}{text}"
        else:
            display = text

        style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
        try:
            p = self.doc.add_paragraph(style=style_name)
        except KeyError:
            p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(display)
        r.font.name = BODY_FONT
        r.font.bold = True
        if level <= 2:
            r.font.size = H1_SIZE if level == 1 else H2_SIZE
            r.font.color.rgb = NAVY
        else:
            r.font.size = H3_SIZE
            r.font.italic = True
        return display

    def paragraph(self, text: str, indent: bool = True, italic: bool = False,
                  center: bool = False, bold: bool = False) -> None:
        p = self.doc.add_paragraph()
        if not indent:
            p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = BODY_SIZE
        r.font.italic = italic
        r.font.bold = bold

    def figure(self, path: str | Path, caption_text: str,
               width: float = 6.5) -> str:
        self._fig_num += 1
        label = f"Fig. {self._fig_num}"
        path = Path(path)
        if path.suffix.lower() == ".pdf":
            png = path.with_suffix(".png")
            if png.exists():
                path = png
            else:
                try:
                    import fitz
                    doc = fitz.open(str(path))
                    pix = doc[0].get_pixmap(dpi=200)
                    pix.save(str(png))
                    doc.close()
                    path = png
                except Exception:
                    pass
        if path.exists() and path.suffix.lower() in (".png", ".jpg", ".jpeg"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run()
            run.add_picture(str(path), width=Inches(width))
        else:
            self.paragraph(f"[{label} placeholder: {path.name}]", indent=False, italic=True)
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"{label}. {caption_text}")
        r.font.name = BODY_FONT
        r.font.size = CAP_SIZE
        r.font.italic = True
        return label

    def table(self, headers: list[str], rows: list[list[str]],
              caption_text: str = "") -> str:
        self._tbl_num += 1
        label = f"Table {self._tbl_num}"

        # Caption ABOVE the table
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}.")
        r.font.name = BODY_FONT
        r.font.size = CAP_SIZE
        r.font.bold = True
        if caption_text:
            r2 = p.add_run(f" {caption_text}")
            r2.font.name = BODY_FONT
            r2.font.size = CAP_SIZE
            r2.font.italic = True

        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_borders(tbl)
        _add_rule(tbl.rows[0], "top", "003366", "12")
        _add_rule(tbl.rows[0], "bottom", "003366", "6")

        for j, h in enumerate(headers):
            cell = tbl.rows[0].cells[j]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(h)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.bold = True
            _shade(cell, HEADER_FILL)

        for i, row_data in enumerate(rows):
            for j, val in enumerate(row_data):
                cell = tbl.rows[i + 1].cells[j]
                cell.text = ""
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(val))
                r.font.name = BODY_FONT
                r.font.size = Pt(9)
                if i % 2 == 1:
                    _shade(cell, ALT_ROW_FILL)

        _add_rule(tbl.rows[-1], "bottom", "003366", "8")
        self.doc.add_paragraph()
        return label

    def equation(self, latex: str) -> str:
        self._eq_num += 1
        clean = latex.strip()
        if not clean:
            return f"Eq. ({self._eq_num})"

        try:
            from equation_handler import add_equation
            success = add_equation(self.doc, clean, self._eq_num)
            if success:
                return f"Eq. ({self._eq_num})"
        except ImportError:
            pass

        # Fallback: italic text with equation number
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"  {clean}    ({self._eq_num})")
        r.font.name = BODY_FONT
        r.font.size = BODY_SIZE
        r.font.italic = True
        return f"Eq. ({self._eq_num})"

    def nomenclature(self, items: list[tuple[str, str, str]]) -> None:
        """Add a nomenclature table. Items: [(symbol, units, meaning), ...]"""
        self.heading("Nomenclature", level=1, numbered=False)
        headers = ["Symbol", "Units", "Description"]
        rows = [[s, u, m] for s, u, m in items]
        self.table(headers, rows, caption_text="")

    def get_stats(self) -> dict:
        """Return word count, paragraph count, figure count, table count, equation count."""
        word_count = 0
        para_count = 0
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text:
                para_count += 1
                word_count += len(text.split())
        return {
            "words": word_count,
            "paragraphs": para_count,
            "figures": self._fig_num,
            "tables": self._tbl_num,
            "equations": self._eq_num,
            "pages_estimate": max(1, word_count // 300),
        }

    def add_stats_footer(self, limits: dict | None = None) -> None:
        """Add a word/page count summary at the end. Optionally check against limits."""
        stats = self.get_stats()
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        parts = [
            f"Words: {stats['words']:,}",
            f"Paragraphs: {stats['paragraphs']}",
            f"Figures: {stats['figures']}",
            f"Tables: {stats['tables']}",
            f"Equations: {stats['equations']}",
            f"Est. pages: {stats['pages_estimate']}",
        ]
        r = p.add_run(" | ".join(parts))
        r.font.name = BODY_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        if limits:
            violations = []
            if "max_words" in limits and stats["words"] > limits["max_words"]:
                violations.append(f"OVER word limit: {stats['words']:,} > {limits['max_words']:,}")
            if "max_pages" in limits and stats["pages_estimate"] > limits["max_pages"]:
                violations.append(f"OVER page limit: ~{stats['pages_estimate']} > {limits['max_pages']}")
            if "max_figures" in limits and stats["figures"] > limits["max_figures"]:
                violations.append(f"OVER figure limit: {stats['figures']} > {limits['max_figures']}")
            if violations:
                p2 = self.doc.add_paragraph()
                p2.paragraph_format.first_line_indent = Pt(0)
                r2 = p2.add_run("WARNING: " + "; ".join(violations))
                r2.font.name = BODY_FONT
                r2.font.size = Pt(9)
                r2.font.bold = True
                r2.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    def credit(self, text: str):
        self.heading("CRediT Authorship Contribution Statement", level=1, numbered=False)
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Bold author names, regular contribution text
        import re
        parts = re.split(r'(Kyeong-Sun Kim|Sung-Ryul Kim)', text)
        for part in parts:
            if part in ("Kyeong-Sun Kim", "Sung-Ryul Kim"):
                r = p.add_run(part)
                r.font.name = BODY_FONT
                r.font.size = BODY_SIZE
                r.font.bold = True
            else:
                r = p.add_run(part)
                r.font.name = BODY_FONT
                r.font.size = BODY_SIZE

    def coi(self, text: str):
        self.heading("Declaration of Competing Interests", level=1, numbered=False)
        self.paragraph(text, indent=False)

    def ai_disclosure(self, text: str):
        self.heading("Declaration of Generative AI and AI-Assisted Technologies", level=1, numbered=False)
        self.paragraph(text, indent=False)

    def acknowledgements(self, text: str):
        self.heading("Acknowledgements", level=1, numbered=False)
        self.paragraph(text, indent=False)

    def references_list(self, refs: list[str]):
        self.heading("References", level=1, numbered=False)
        for ref in refs:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(ref)
            r.font.name = BODY_FONT
            r.font.size = Pt(10)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        self.doc.save(str(path))
        log.info("saved %s (%d bytes)", path.name, path.stat().st_size)
        return path
