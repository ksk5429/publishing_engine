"""
LaTeX → native Word OMML equation handler.

Two methods available (in priority order):
  1. markdocx's math_renderer: LaTeX → MathML → OMML (pure Python, no external deps)
  2. pandoc subprocess: LaTeX → .docx → extract OMML (fallback)

Both produce EDITABLE native Word equations — not images.
"""
from __future__ import annotations

import copy
import logging
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

sys.path.insert(0, str(Path(__file__).parent))

log = logging.getLogger("equation_handler")

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _try_markdocx(latex: str) -> etree._Element | None:
    """Use markdocx's MathML→OMML converter (pure Python)."""
    try:
        from math_renderer import latex_to_omml_para, latex_to_omml
        result = latex_to_omml_para(latex)
        if result is not None:
            return result
        result = latex_to_omml(latex)
        return result
    except Exception as e:
        log.debug("markdocx math_renderer failed: %s", e)
        return None


def _try_pandoc(latex: str) -> etree._Element | None:
    """Use pandoc subprocess to convert LaTeX → OMML."""
    pandoc_paths = [
        Path(r"C:/Users/geolab/AppData/Local/Pandoc/pandoc.exe"),
        Path(r"C:/Program Files/Pandoc/pandoc.exe"),
    ]
    pandoc = None
    for p in pandoc_paths:
        if p.exists():
            pandoc = str(p)
            break
    if pandoc is None:
        return None

    with tempfile.TemporaryDirectory() as tmp:
        md_path = Path(tmp) / "eq.md"
        docx_path = Path(tmp) / "eq.docx"
        md_path.write_text(f"$${latex}$$", encoding="utf-8")
        result = subprocess.run(
            [pandoc, str(md_path), "-o", str(docx_path)],
            capture_output=True, text=True, timeout=15,
        )
        if result.returncode != 0 or not docx_path.exists():
            return None
        try:
            doc = Document(str(docx_path))
        except Exception:
            return None
        ns = {"m": OMML_NS}
        for p in doc.paragraphs:
            found = p._p.findall(f".//{{{OMML_NS}}}oMathPara")
            if found:
                return copy.deepcopy(found[0])
            found = p._p.findall(f".//{{{OMML_NS}}}oMath")
            if found:
                return copy.deepcopy(found[0])
    return None


def add_equation(doc: Document, latex: str, eq_num: int) -> bool:
    """Add a display equation with native OMML and equation number.
    Returns True if native rendering succeeded."""
    clean = latex.strip()
    if not clean:
        return False

    # Try markdocx first (fast, pure Python), then pandoc (subprocess)
    omml = _try_markdocx(clean)
    if omml is None:
        log.debug("markdocx failed for eq %d, trying pandoc...", eq_num)
        omml = _try_pandoc(clean)

    if omml is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p._p.append(omml)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.first_line_indent = Pt(0)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        r = p2.add_run(f"({eq_num})")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        return True

    # Fallback: italic text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(f"  {clean}    ({eq_num})")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.italic = True
    log.warning("equation %d: fallback italic text", eq_num)
    return False
