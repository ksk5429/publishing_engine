"""
Generate a Response to Reviewers DOCX.

Reads a Reviewers_Comments.txt and a responses YAML/dict, produces a
publication-quality point-by-point response document with:
  - Reviewer comment (grey box)
  - Author response (blue text)
  - Action taken (location in revised manuscript)

Usage:
    python _shared/render_response.py _shared/sample
"""
from __future__ import annotations

import argparse
import logging
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

PAPERS_DIR = Path(__file__).parent.parent  # repo root
BODY_FONT = "Times New Roman"
NAVY = RGBColor(0x00, 0x33, 0x66)
BLUE = RGBColor(0x15, 0x65, 0xC0)
GREY_BG = "F0F0F0"
RED = RGBColor(0xC6, 0x28, 0x28)

logging.basicConfig(level=logging.INFO, format="[%(asctime)s] %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("response")


def _shade(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _remove_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    if tbl_pr is None:
        return
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _add_left_border(cell, color="1565C0", sz="18"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:color"), color)
    borders.append(left)
    tc_pr.append(borders)


def parse_comments(txt_path: Path) -> list[dict]:
    """Parse reviewer comments from a text file into structured items."""
    text = txt_path.read_text(encoding="utf-8")
    items = []
    current_reviewer = ""

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Detect reviewer header
        m = re.match(r"Reviewer\s*#?(\d+)", line, re.IGNORECASE)
        if m:
            current_reviewer = f"Reviewer #{m.group(1)}"
            continue
        # Detect comment ID
        m = re.match(r"(R\d+\.\d+)[:\s]+(.*)", line)
        if m:
            items.append({
                "reviewer": current_reviewer,
                "id": m.group(1),
                "comment": m.group(2).strip(),
            })
        elif items and not line.startswith("="):
            items[-1]["comment"] += " " + line

    return items


def build_response(paper_dir: Path, responses: dict[str, str] | None = None) -> Path:
    """Build the response-to-reviewers DOCX."""
    comments_path = paper_dir / "Reviewers_Comments.txt"
    if not comments_path.exists():
        raise FileNotFoundError(f"No Reviewers_Comments.txt in {paper_dir}")

    items = parse_comments(comments_path)
    log.info("parsed %d reviewer comments from %s", len(items), comments_path.name)

    if responses is None:
        responses = {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("Response to Reviewers")
    r.font.name = BODY_FONT
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("We thank the reviewers for their constructive comments. "
                   "Each comment is addressed below with our response and the "
                   "specific changes made in the revised manuscript.")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.italic = True

    doc.add_paragraph()

    # Summary table
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("Summary of Changes")
    r.font.name = BODY_FONT
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = NAVY

    tbl = doc.add_table(rows=1 + len(items), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Comment (summary)", "Status", "Section"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = BODY_FONT
                r.font.size = Pt(8)
                r.font.bold = True
        _shade(cell, "1F3864")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, item in enumerate(items):
        row = tbl.rows[i + 1]
        resp = responses.get(item["id"], {})
        vals = [
            item["id"],
            item["comment"][:60] + ("..." if len(item["comment"]) > 60 else ""),
            resp.get("status", "Addressed"),
            resp.get("section", "—"),
        ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = BODY_FONT
                    r.font.size = Pt(8)
            if i % 2 == 1:
                _shade(cell, "F2F6FC")

    doc.add_page_break()

    # Point-by-point responses
    current_reviewer = ""
    for item in items:
        if item["reviewer"] != current_reviewer:
            current_reviewer = item["reviewer"]
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(18)
            r = p.add_run(current_reviewer)
            r.font.name = BODY_FONT
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = NAVY

        # Comment box (grey background, single-cell table)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(f"{item['id']} — Reviewer Comment:")
        r.font.name = BODY_FONT
        r.font.size = Pt(10)
        r.font.bold = True

        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        cell.text = item["comment"]
        _shade(cell, GREY_BG)
        for p in cell.paragraphs:
            p.paragraph_format.first_line_indent = Pt(0)
            for r in p.runs:
                r.font.name = BODY_FONT
                r.font.size = Pt(10)
                r.font.italic = True

        # Response (blue left-border)
        resp = responses.get(item["id"], {})
        resp_text = resp.get("response", "[Author response to be added]")

        tbl2 = doc.add_table(rows=1, cols=1)
        cell2 = tbl2.rows[0].cells[0]
        _add_left_border(cell2, "1565C0", "18")
        p = cell2.paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run("Author Response: ")
        r.font.name = BODY_FONT
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = BLUE
        r = p.add_run(resp_text)
        r.font.name = BODY_FONT
        r.font.size = Pt(10)
        r.font.color.rgb = BLUE

        # Action taken
        action = resp.get("action", "")
        if action:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run("Action: ")
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RED
            r = p.add_run(action)
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.color.rgb = RED

    out_dir = paper_dir / "_output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "response_to_reviewers.docx"
    doc.save(str(out_path))
    log.info("wrote %s (%d bytes)", out_path, out_path.stat().st_size)
    return out_path


# Sample responses for the demo
SAMPLE_RESPONSES = {
    "R1.1": {
        "response": "We have shortened the abstract to 195 words, focusing on the three key findings: (1) coherence as the only sign-consistent feature, (2) the pi6 negative result, and (3) the cross-site transfer bound.",
        "action": "Abstract revised (Page 1)",
        "status": "Revised",
        "section": "Abstract",
    },
    "R1.2": {
        "response": "We have added a clarification that a = -0.167 represents the fractional frequency reduction per unit (S/D)^b, and b = 1.47 is the nonlinearity exponent. The calibration is valid for 0 <= S/D <= 0.5, which covers the range tested in the centrifuge programme.",
        "action": "Paragraph added after Eq. (1) in Section 2",
        "status": "Revised",
        "section": "§2",
    },
    "R1.3": {
        "response": "Table 1 now includes the centrifuge g-level (70g) and model scale factor (1:70) as additional columns.",
        "action": "Table 1 updated",
        "status": "Revised",
        "section": "§3, Table 1",
    },
    "R1.4": {
        "response": "We have added a paragraph in the Discussion (Section 6) explicitly addressing the parked-state limitation. Under operational loading, rotor-induced harmonics may mask or confound the coherence feature. Extension to operational states requires excitation-aware normalisation, which is identified as future work.",
        "action": "New paragraph in Section 6",
        "status": "Revised",
        "section": "§6",
    },
    "R2.1": {
        "response": "We have added a comparison with PCA (3-component removal) on the same dataset. The state-function method achieves <1% FAR versus ~3% for PCA, while detection latency is identical for the January 2024 event.",
        "action": "New subsection 5.4 added with comparison table",
        "status": "Revised",
        "section": "§5.4 (new)",
    },
    "R2.2": {
        "response": "The +55.8% shift in mid-elevation strain-acceleration coherence is consistent with a loss of lateral restraint at the scoured bucket, which alters the strain distribution along the tower at the first natural frequency. This mechanism is confirmed by the centrifuge data where T4 shows a +12% coherence increase at S/D = 0.58.",
        "action": "Explanation added to Section 5.3",
        "status": "Revised",
        "section": "§5.3",
    },
    "R2.3": {
        "response": "We have added citations for the frequency suppression claim (Jalbi and Bhattacharya, 2018) and the 0.05D detection resolution claim (Weil et al., 2023). Both are now properly referenced in Section 1.",
        "action": "Two citations added to Section 1, paragraphs 1 and 2",
        "status": "Revised",
        "section": "§1",
    },
    "R2.4": {
        "response": "A Data and Code Availability statement has been added. The Op3 framework (v1.0.0-rc2) is available on PyPI. All analysis scripts are provided in the supplementary material. The centrifuge feature CSV and field feature matrix will be deposited on Zenodo upon acceptance.",
        "action": "New Data and Code Availability section added before References",
        "status": "Revised",
        "section": "Data Availability",
    },
}


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("paper", help="Paper directory")
    args = parser.parse_args()
    paper_dir = PAPERS_DIR / args.paper
    build_response(paper_dir, SAMPLE_RESPONSES)


if __name__ == "__main__":
    main()
